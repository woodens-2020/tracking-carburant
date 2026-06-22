"""
Trois rendeurs indépendants consommant la même structure de payload.

  render_pdf(payload, narrative, charts)  -> bytes
  render_docx(payload, narrative, charts) -> bytes
  render_xlsx(payload, narrative, charts) -> bytes

Aucun calcul métier ici — tout vient du payload produit par rapport_service.py.
"""
from __future__ import annotations

from io import BytesIO
from typing import Optional


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS COMMUNS
# ══════════════════════════════════════════════════════════════════════════════

def _fmt_g(v: float) -> str:
    return f"{v:,.2f} G" if v is not None else "—"

def _fmt_gal(v: float) -> str:
    return f"{v:,.3f} gal" if v is not None else "—"

def _fmt_pct(v: Optional[float]) -> str:
    if v is None:
        return "—"
    signe = "+" if v >= 0 else ""
    return f"{signe}{v:.1f} %"


# ══════════════════════════════════════════════════════════════════════════════
# 1. PDF  (reportlab)
# ══════════════════════════════════════════════════════════════════════════════

def render_pdf(payload: dict, narrative: dict, charts: dict[str, bytes]) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        HRFlowable, Image, KeepTogether, PageBreak, Paragraph, SimpleDocTemplate,
        Spacer, Table, TableStyle,
    )

    stats = payload["stats"]
    rentab = payload["rentab"]
    total_rentab = rentab.get("total", {})
    anomalies = payload["anomalies"]
    stocks = payload["stocks"]

    # ── Couleurs ──────────────────────────────────────────────────
    NAVY    = colors.HexColor("#0f1e35")
    AMBER   = colors.HexColor("#f7a93b")
    TEAL    = colors.HexColor("#3fb6a8")
    TEAL_L  = colors.HexColor("#e0f5f3")
    AMBER_L = colors.HexColor("#fef3d7")
    RED     = colors.HexColor("#dc2626")
    RED_L   = colors.HexColor("#fee2e2")
    GREY    = colors.HexColor("#64748b")
    GREY_L  = colors.HexColor("#f8fafc")
    GREY_B  = colors.HexColor("#e2e8f0")
    WHITE   = colors.white

    def S(name, **kw):
        base = getSampleStyleSheet()
        parent = base.get(name, base["Normal"])
        return ParagraphStyle(name + "_custom", parent=parent, **kw)

    sH1   = S("Heading1", fontSize=16, textColor=NAVY, spaceAfter=4,
               fontName="Helvetica-Bold", leading=20)
    sH2   = S("Heading2", fontSize=12, textColor=NAVY, spaceAfter=3,
               spaceBefore=10, fontName="Helvetica-Bold")
    sH3   = S("Heading3", fontSize=10, textColor=TEAL, spaceAfter=2,
               spaceBefore=6, fontName="Helvetica-Bold")
    sBody = S("Normal", fontSize=9, textColor=colors.HexColor("#1e293b"),
              leading=13, spaceAfter=3)
    sJust = S("Just", fontSize=9, textColor=colors.HexColor("#1e293b"),
              leading=13, alignment=TA_JUSTIFY, spaceAfter=3)
    sMut  = S("Muted", fontSize=8, textColor=GREY, leading=11)
    sCtr  = S("Center", fontSize=9, alignment=TA_CENTER)

    def _ts_base():
        return [
            ("FONTSIZE",      (0, 0), (-1, -1), 8.5),
            ("TOPPADDING",    (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING",   (0, 0), (-1, -1), 6),
            ("GRID",          (0, 0), (-1, -1), 0.3, GREY_B),
            ("VALIGN",        (0, 0), (-1, -1), "TOP"),
        ]

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=1.8 * cm, rightMargin=1.8 * cm,
        topMargin=2.0 * cm, bottomMargin=2.0 * cm,
        title=f"Rapport {payload['date_debut']} — {payload['date_fin']}",
        author="PétroSync",
    )
    story = []

    def hr(color=GREY_B, thick=0.5):
        story.append(HRFlowable(width="100%", thickness=thick, color=color, spaceAfter=6))

    def img_from_bytes(data: bytes, max_w=14*cm, max_h=8*cm):
        buf_img = BytesIO(data)
        img = Image(buf_img)
        ratio = img.imageWidth / img.imageHeight
        w = min(max_w, max_h * ratio)
        h = w / ratio
        img._width = w
        img._height = h
        img.drawWidth = w
        img.drawHeight = h
        return img

    # ─── PAGE DE GARDE ────────────────────────────────────────────
    story.append(Spacer(1, 2.5 * cm))
    story.append(Paragraph("RAPPORT DE GESTION", S("cov0", fontSize=10, textColor=AMBER,
                            fontName="Helvetica-Bold", alignment=TA_CENTER)))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(payload["station_nom"].upper(), S("cov1", fontSize=20,
                            textColor=NAVY, fontName="Helvetica-Bold", alignment=TA_CENTER)))
    story.append(Spacer(1, 0.5 * cm))

    periode_str = (
        payload["date_debut"] if payload["date_debut"] == payload["date_fin"]
        else f"{payload['date_debut']}  →  {payload['date_fin']}"
    )
    story.append(Paragraph(periode_str, S("per", fontSize=13, textColor=TEAL,
                            alignment=TA_CENTER, fontName="Helvetica-Bold")))
    story.append(Spacer(1, 0.8 * cm))
    hr(AMBER, 2)

    meta = [
        ["Généré le", payload["date_generation"]],
        ["Période", f"{payload['nb_jours']} jour(s) — {payload['date_debut']} au {payload['date_fin']}"],
        ["Relevés analysés", str(stats["nb_releves"])],
        ["Jours actifs", str(stats["nb_jours_couverts"])],
        ["Anomalies détectées", str(len(anomalies))],
    ]
    mt = Table(meta, colWidths=[4 * cm, 12 * cm])
    mt.setStyle(TableStyle([
        ("FONTNAME",     (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE",     (0, 0), (-1, -1), 8.5),
        ("TEXTCOLOR",    (0, 0), (0, -1), GREY),
        ("TEXTCOLOR",    (1, 0), (1, -1), NAVY),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [GREY_L, WHITE]),
        ("GRID",         (0, 0), (-1, -1), 0.3, GREY_B),
        ("TOPPADDING",   (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 5),
        ("LEFTPADDING",  (0, 0), (-1, -1), 8),
    ]))
    story.append(mt)
    story.append(PageBreak())

    # ─── 1. RÉSUMÉ EXÉCUTIF ───────────────────────────────────────
    story.append(Paragraph("1. Résumé Exécutif", sH1))
    hr(AMBER, 1.5)

    story.append(Paragraph(narrative["intro_kpis"], sJust))
    story.append(Spacer(1, 0.3 * cm))

    # KPI banner (tableau 4 colonnes)
    kpi_data = [
        [Paragraph("<b>Volume vendu</b>", sCtr),
         Paragraph("<b>Revenu total</b>", sCtr),
         Paragraph("<b>Variation vs N-1</b>", sCtr),
         Paragraph("<b>Anomalies</b>", sCtr)],
        [Paragraph(f"<b>{_fmt_gal(stats['total_quantite'])}</b>",
                   S("kv", fontSize=13, textColor=TEAL, fontName="Helvetica-Bold", alignment=TA_CENTER)),
         Paragraph(f"<b>{_fmt_g(stats['total_montant'])}</b>",
                   S("km", fontSize=13, textColor=AMBER, fontName="Helvetica-Bold", alignment=TA_CENTER)),
         Paragraph(f"<b>{_fmt_pct(narrative['var_pct'])}</b>",
                   S("kvar", fontSize=13,
                     textColor=RED if (narrative["var_pct"] or 0) < 0 else TEAL,
                     fontName="Helvetica-Bold", alignment=TA_CENTER)),
         Paragraph(f"<b>{len(anomalies)}</b>",
                   S("kan", fontSize=13,
                     textColor=RED if anomalies else TEAL,
                     fontName="Helvetica-Bold", alignment=TA_CENTER))],
    ]
    kt = Table(kpi_data, colWidths=[4 * cm] * 4)
    kt.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, 0), NAVY),
        ("TEXTCOLOR",     (0, 0), (-1, 0), WHITE),
        ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",      (0, 0), (-1, 0), 8),
        ("ROWBACKGROUNDS",(0, 1), (-1, -1), [GREY_L]),
        ("GRID",          (0, 0), (-1, -1), 0.3, GREY_B),
        ("TOPPADDING",    (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("ALIGN",         (0, 0), (-1, -1), "CENTER"),
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(kt)
    story.append(PageBreak())

    # ─── 2. ANALYSE DES VENTES ───────────────────────────────────
    story.append(Paragraph("2. Analyse des Ventes", sH1))
    hr(TEAL, 1.5)

    story.append(Paragraph("2.1 Répartition par produit", sH2))
    story.append(Paragraph(narrative["ventes_text"], sJust))

    if stats["par_produit"]:
        hdr_v = [["Produit", "Volume (gal)", "Revenu (G)", "% du revenu"]]
        rows_v = []
        for nom, d in sorted(stats["par_produit"].items(),
                              key=lambda x: x[1]["montant"], reverse=True):
            pct = (round(d["montant"] / stats["total_montant"] * 100, 1)
                   if stats["total_montant"] > 0 else 0.0)
            rows_v.append([nom, f"{d['quantite']:,.3f}", f"{d['montant']:,.2f}", f"{pct:.1f} %"])
        tv = Table(hdr_v + rows_v, colWidths=[5*cm, 4*cm, 4*cm, 3*cm])
        tv.setStyle(TableStyle(_ts_base() + [
            ("BACKGROUND",    (0, 0), (-1, 0), NAVY),
            ("TEXTCOLOR",     (0, 0), (-1, 0), WHITE),
            ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
            ("ROWBACKGROUNDS",(0, 1), (-1, -1), [GREY_L, WHITE]),
            ("ALIGN",         (1, 0), (-1, -1), "RIGHT"),
        ]))
        story.append(tv)
        story.append(Spacer(1, 0.3 * cm))

    if "par_produit" in charts:
        story.append(img_from_bytes(charts["par_produit"], max_w=9*cm, max_h=7*cm))
        story.append(Spacer(1, 0.3 * cm))

    story.append(Paragraph("2.2 Répartition par pompe", sH2))
    story.append(Paragraph(narrative["pompes_text"], sJust))

    if stats["par_pompe"]:
        hdr_p = [["Pompe", "Produit", "Volume (gal)", "Revenu (G)"]]
        rows_p = [
            [nom, d["produit"], f"{d['quantite']:,.3f}", f"{d['montant']:,.2f}"]
            for nom, d in sorted(stats["par_pompe"].items(),
                                 key=lambda x: x[1]["montant"], reverse=True)
        ]
        tp = Table(hdr_p + rows_p, colWidths=[4*cm, 4*cm, 4*cm, 4*cm])
        tp.setStyle(TableStyle(_ts_base() + [
            ("BACKGROUND",    (0, 0), (-1, 0), NAVY),
            ("TEXTCOLOR",     (0, 0), (-1, 0), WHITE),
            ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
            ("ROWBACKGROUNDS",(0, 1), (-1, -1), [GREY_L, WHITE]),
            ("ALIGN",         (2, 0), (-1, -1), "RIGHT"),
        ]))
        story.append(tp)
        story.append(Spacer(1, 0.3 * cm))

    if "par_pompe" in charts:
        story.append(img_from_bytes(charts["par_pompe"], max_w=13*cm, max_h=6*cm))

    story.append(Paragraph("2.3 Répartition par période", sH2))
    story.append(Paragraph(narrative["periode_text"], sJust))

    if "ventes_jours" in charts:
        story.append(Paragraph("2.4 Évolution journalière", sH2))
        story.append(img_from_bytes(charts["ventes_jours"], max_w=16*cm, max_h=6*cm))

    story.append(PageBreak())

    # ─── 3. ANOMALIES ────────────────────────────────────────────
    story.append(Paragraph("3. Analyse des Anomalies", sH1))
    hr(RED, 1.5)
    story.append(Paragraph(narrative["anom_text"], sJust))
    story.append(Spacer(1, 0.2 * cm))

    if anomalies:
        if "anomalies" in charts:
            story.append(img_from_bytes(charts["anomalies"], max_w=12*cm, max_h=5*cm))
            story.append(Spacer(1, 0.3 * cm))

        hdr_a = [["Type", "Gravité", "Pompe", "Date", "Période", "Message"]]
        rows_a = [
            [
                a.get("type", "—"),
                a.get("gravite", "—"),
                a.get("pompe_nom", a.get("produit_nom", "—")),
                a.get("date", "—"),
                a.get("periode", "—"),
                (a.get("message", "") or "")[:80],
            ]
            for a in anomalies[:30]
        ]
        ta = Table(hdr_a + rows_a, colWidths=[3.5*cm, 2*cm, 2.5*cm, 2*cm, 2*cm, 4*cm])
        ts_a = TableStyle(_ts_base() + [
            ("BACKGROUND",    (0, 0), (-1, 0), NAVY),
            ("TEXTCOLOR",     (0, 0), (-1, 0), WHITE),
            ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
            ("ROWBACKGROUNDS",(0, 1), (-1, -1), [GREY_L, WHITE]),
            ("FONTSIZE",      (0, 0), (-1, -1), 7.5),
        ])
        for i, a in enumerate(anomalies[:30], 1):
            if a.get("gravite") == "erreur":
                ts_a.add("BACKGROUND", (1, i), (1, i), RED_L)
                ts_a.add("TEXTCOLOR",  (1, i), (1, i), RED)
                ts_a.add("FONTNAME",   (1, i), (1, i), "Helvetica-Bold")
        ta.setStyle(ts_a)
        story.append(ta)
    else:
        story.append(Paragraph("✓ Aucune anomalie détectée.", S("ok", fontSize=10,
                                textColor=TEAL, fontName="Helvetica-Bold")))

    story.append(PageBreak())

    # ─── 4. STOCK & RENTABILITÉ ───────────────────────────────────
    story.append(Paragraph("4. Stock et Rentabilité", sH1))
    hr(TEAL, 1.5)

    story.append(Paragraph("4.1 Niveaux de stock actuels", sH2))
    story.append(Paragraph(narrative["stock_text"], sJust))

    if stocks:
        hdr_s = [["Produit", "Stock (gal)", "Livrés (gal)", "Vendus (gal)", "Jours rest.", "Alerte"]]
        rows_s = [
            [
                s["produit_nom"],
                f"{s['gallons_restants']:,.3f}",
                f"{s['gallons_livres']:,.3f}",
                f"{s['gallons_vendus']:,.3f}",
                f"{s['jours_de_stock']:.1f}" if s["jours_de_stock"] is not None else "—",
                "⚠ BAS" if s["alerte_bas"] else "OK",
            ]
            for s in stocks
        ]
        ts_s = Table(hdr_s + rows_s, colWidths=[3.5*cm, 2.8*cm, 2.8*cm, 2.8*cm, 2.3*cm, 1.8*cm])
        ts_sty = TableStyle(_ts_base() + [
            ("BACKGROUND",    (0, 0), (-1, 0), NAVY),
            ("TEXTCOLOR",     (0, 0), (-1, 0), WHITE),
            ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
            ("ROWBACKGROUNDS",(0, 1), (-1, -1), [GREY_L, WHITE]),
            ("ALIGN",         (1, 0), (-1, -1), "RIGHT"),
        ])
        for i, s in enumerate(stocks, 1):
            if s["alerte_bas"]:
                ts_sty.add("BACKGROUND", (-1, i), (-1, i), RED_L)
                ts_sty.add("TEXTCOLOR",  (-1, i), (-1, i), RED)
                ts_sty.add("FONTNAME",   (-1, i), (-1, i), "Helvetica-Bold")
        ts_s.setStyle(ts_sty)
        story.append(ts_s)
        story.append(Spacer(1, 0.3 * cm))

    story.append(Paragraph("4.2 Rentabilité (Coût Moyen Pondéré)", sH2))
    story.append(Paragraph(narrative["rentab_text"], sJust))

    if rentab.get("produits"):
        hdr_r = [["Produit", "Vendus (gal)", "Revenu (G)", "COGS (G)", "Bénéfice (G)", "Marge %"]]
        rows_r = [
            [
                p["produit_nom"],
                f"{p['gallons_vendus']:,.3f}",
                f"{p['revenu_total']:,.2f}",
                f"{p['cogs_total']:,.2f}" if p["cogs_total"] is not None else "—",
                f"{p['benefice']:,.2f}" if p["benefice"] is not None else "—",
                f"{p['marge_pct']:.2f} %" if p["marge_pct"] is not None else "—",
            ]
            for p in rentab["produits"]
        ]
        # Ligne total
        t = rentab.get("total", {})
        rows_r.append([
            "TOTAL",
            f"{t.get('gallons_vendus', 0):,.3f}",
            f"{t.get('revenu_total', 0):,.2f}",
            f"{t.get('cogs_total', 0):,.2f}" if t.get("cogs_total") is not None else "—",
            f"{t.get('benefice', 0):,.2f}" if t.get("benefice") is not None else "—",
            f"{t.get('marge_pct', 0):.2f} %" if t.get("marge_pct") is not None else "—",
        ])
        tr = Table(hdr_r + rows_r, colWidths=[3*cm, 2.8*cm, 2.8*cm, 2.8*cm, 2.8*cm, 1.8*cm])
        tr.setStyle(TableStyle(_ts_base() + [
            ("BACKGROUND",    (0, 0), (-1, 0), NAVY),
            ("TEXTCOLOR",     (0, 0), (-1, 0), WHITE),
            ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
            ("ROWBACKGROUNDS",(0, 1), (-1, -2), [GREY_L, WHITE]),
            ("BACKGROUND",    (0, -1), (-1, -1), AMBER_L),
            ("FONTNAME",      (0, -1), (-1, -1), "Helvetica-Bold"),
            ("ALIGN",         (1, 0), (-1, -1), "RIGHT"),
        ]))
        story.append(tr)

    story.append(PageBreak())

    # ─── 5. CONCLUSION & RECOMMANDATIONS ─────────────────────────
    story.append(Paragraph("5. Conclusion", sH1))
    hr(AMBER, 1.5)
    story.append(Paragraph(narrative["conclusion"], sJust))
    story.append(Spacer(1, 0.4 * cm))

    story.append(Paragraph("6. Recommandations", sH1))
    hr(AMBER, 1.5)
    for i, rec in enumerate(narrative["recommandations"], 1):
        story.append(Paragraph(f"{i}. {rec}", sBody))

    doc.build(story)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════════════════════
# 2. WORD / DOCX  (python-docx)
# ══════════════════════════════════════════════════════════════════════════════

def render_docx(payload: dict, narrative: dict, charts: dict[str, bytes]) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
    from docx.oxml import OxmlElement

    stats = payload["stats"]
    rentab = payload["rentab"]
    stocks = payload["stocks"]
    anomalies = payload["anomalies"]

    NAVY   = RGBColor(0x0f, 0x1e, 0x35)
    AMBER  = RGBColor(0xf7, 0xa9, 0x3b)
    TEAL   = RGBColor(0x3f, 0xb6, 0xa8)
    RED    = RGBColor(0xdc, 0x26, 0x26)
    GREY   = RGBColor(0x64, 0x74, 0x8b)

    doc = Document()

    # ── Marges ────────────────────────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    def set_color(run, color: RGBColor):
        run.font.color.rgb = color

    def heading(text: str, level: int = 1, color: RGBColor = NAVY):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            set_color(run, color)
        return h

    def body_text(text: str):
        p = doc.add_paragraph(text)
        p.style.font.size = Pt(10)
        return p

    def add_img(png_bytes: bytes, width_cm: float = 14):
        buf = BytesIO(png_bytes)
        doc.add_picture(buf, width=Cm(width_cm))

    def add_table(headers: list[str], rows: list[list[str]]):
        t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        t.style = "Table Grid"
        hdr_row = t.rows[0]
        for i, h in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                run.font.size = Pt(9)
            # Fond bleu marine pour l'en-tête
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "0F1E35")
            tcPr.append(shd)
        for r_i, row_data in enumerate(rows, 1):
            row = t.rows[r_i]
            bg = "F8FAFC" if r_i % 2 == 0 else "FFFFFF"
            for c_i, val in enumerate(row_data):
                cell = row.cells[c_i]
                cell.text = str(val)
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), bg)
                tcPr.append(shd)
        return t

    # ─── Page de titre ────────────────────────────────────────────
    t = doc.add_heading("RAPPORT DE GESTION", 0)
    for run in t.runs:
        set_color(run, NAVY)
        run.font.size = Pt(20)

    sub = doc.add_paragraph(payload["station_nom"])
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.bold = True
        run.font.size = Pt(14)
        set_color(run, TEAL)

    period = doc.add_paragraph(f"{payload['date_debut']}  →  {payload['date_fin']}")
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in period.runs:
        run.font.size = Pt(12)
        set_color(run, GREY)

    doc.add_paragraph(f"Généré le {payload['date_generation']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ─── 1. Résumé exécutif ───────────────────────────────────────
    heading("1. Résumé Exécutif", 1)
    body_text(narrative["intro_kpis"])
    doc.add_paragraph()

    add_table(
        ["Indicateur", "Valeur"],
        [
            ["Volume vendu", _fmt_gal(stats["total_quantite"])],
            ["Revenu total", _fmt_g(stats["total_montant"])],
            ["Relevés analysés", str(stats["nb_releves"])],
            ["Jours actifs", str(stats["nb_jours_couverts"])],
            ["Variation vs période préc.", _fmt_pct(narrative["var_pct"])],
            ["Anomalies détectées", str(len(anomalies))],
        ],
    )
    doc.add_page_break()

    # ─── 2. Ventes ────────────────────────────────────────────────
    heading("2. Analyse des Ventes", 1)

    heading("2.1 Par produit", 2, TEAL)
    body_text(narrative["ventes_text"])
    if stats["par_produit"]:
        rows_v = [
            [
                nom,
                f"{d['quantite']:,.3f}",
                f"{d['montant']:,.2f}",
                f"{round(d['montant']/stats['total_montant']*100,1):.1f} %"
                if stats["total_montant"] > 0 else "—",
            ]
            for nom, d in sorted(stats["par_produit"].items(),
                                  key=lambda x: x[1]["montant"], reverse=True)
        ]
        add_table(["Produit", "Volume (gal)", "Revenu (G)", "% revenu"], rows_v)
    if "par_produit" in charts:
        doc.add_paragraph()
        add_img(charts["par_produit"], 10)

    heading("2.2 Par pompe", 2, TEAL)
    body_text(narrative["pompes_text"])
    if stats["par_pompe"]:
        rows_p = [
            [nom, d["produit"], f"{d['quantite']:,.3f}", f"{d['montant']:,.2f}"]
            for nom, d in sorted(stats["par_pompe"].items(),
                                  key=lambda x: x[1]["montant"], reverse=True)
        ]
        add_table(["Pompe", "Produit", "Volume (gal)", "Revenu (G)"], rows_p)
    if "par_pompe" in charts:
        doc.add_paragraph()
        add_img(charts["par_pompe"], 13)

    heading("2.3 Par période", 2, TEAL)
    body_text(narrative["periode_text"])

    if "ventes_jours" in charts:
        heading("2.4 Évolution journalière", 2, TEAL)
        add_img(charts["ventes_jours"], 15)

    doc.add_page_break()

    # ─── 3. Anomalies ─────────────────────────────────────────────
    heading("3. Analyse des Anomalies", 1)
    body_text(narrative["anom_text"])
    if "anomalies" in charts:
        doc.add_paragraph()
        add_img(charts["anomalies"], 12)
    if anomalies:
        doc.add_paragraph()
        rows_a = [
            [
                a.get("type", "—"),
                a.get("gravite", "—"),
                a.get("pompe_nom", a.get("produit_nom", "—")),
                a.get("date", "—"),
                a.get("periode", "—"),
                (a.get("message", "") or "")[:100],
            ]
            for a in anomalies[:30]
        ]
        add_table(["Type", "Gravité", "Pompe", "Date", "Période", "Message"], rows_a)

    doc.add_page_break()

    # ─── 4. Stock & Rentabilité ───────────────────────────────────
    heading("4. Stock et Rentabilité", 1)

    heading("4.1 Niveaux de stock", 2, TEAL)
    body_text(narrative["stock_text"])
    if stocks:
        rows_s = [
            [
                s["produit_nom"],
                f"{s['gallons_restants']:,.3f}",
                f"{s['gallons_livres']:,.3f}",
                f"{s['gallons_vendus']:,.3f}",
                f"{s['jours_de_stock']:.1f}" if s["jours_de_stock"] is not None else "—",
                "⚠ BAS" if s["alerte_bas"] else "OK",
            ]
            for s in stocks
        ]
        add_table(["Produit", "Stock (gal)", "Livrés", "Vendus", "Jours", "Alerte"], rows_s)

    heading("4.2 Rentabilité (WAC)", 2, TEAL)
    body_text(narrative["rentab_text"])
    if rentab.get("produits"):
        rows_r = [
            [
                p["produit_nom"],
                f"{p['gallons_vendus']:,.3f}",
                f"{p['revenu_total']:,.2f}",
                f"{p['cogs_total']:,.2f}" if p["cogs_total"] is not None else "—",
                f"{p['benefice']:,.2f}" if p["benefice"] is not None else "—",
                f"{p['marge_pct']:.2f} %" if p["marge_pct"] is not None else "—",
            ]
            for p in rentab["produits"]
        ]
        add_table(["Produit", "Vendus (gal)", "Revenu (G)", "COGS (G)", "Bénéfice (G)", "Marge %"], rows_r)

    doc.add_page_break()

    # ─── 5. Conclusion & Recommandations ─────────────────────────
    heading("5. Conclusion", 1)
    body_text(narrative["conclusion"])

    doc.add_paragraph()
    heading("6. Recommandations", 1)
    for i, rec in enumerate(narrative["recommandations"], 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(rec).font.size = Pt(10)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════════════════════
# 3. EXCEL / XLSX  (openpyxl)
# ══════════════════════════════════════════════════════════════════════════════

def render_xlsx(payload: dict, narrative: dict, charts: dict[str, bytes]) -> bytes:
    from openpyxl import Workbook
    from openpyxl.chart import BarChart, PieChart, Reference
    from openpyxl.chart.series import SeriesLabel
    from openpyxl.drawing.image import Image as XlImage
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    stats = payload["stats"]
    rentab = payload["rentab"]
    stocks = payload["stocks"]
    anomalies = payload["anomalies"]
    serie = payload["serie_jours"]

    # ── Styles ────────────────────────────────────────────────────
    NAVY_H  = "0F1E35"
    AMBER_H = "F7A93B"
    TEAL_H  = "3FB6A8"
    RED_H   = "DC2626"
    REDL_H  = "FEE2E2"
    GREY_H  = "F8FAFC"
    GREYB_H = "E2E8F0"
    WHITE   = "FFFFFF"
    TEAL_L  = "E0F5F3"

    def _fill(hex_: str) -> PatternFill:
        return PatternFill("solid", fgColor=hex_)

    def _font(bold=False, color="1E293B", size=10, italic=False) -> Font:
        return Font(bold=bold, color=color, size=size, italic=italic)

    def _align(h="left", v="center", wrap=False) -> Alignment:
        return Alignment(horizontal=h, vertical=v, wrap_text=wrap)

    def _border() -> Border:
        s = Side(style="thin", color=GREYB_H)
        return Border(left=s, right=s, top=s, bottom=s)

    def _hdr(ws, row: int, data: list, bg=NAVY_H, fg=WHITE, size=10):
        for col, val in enumerate(data, 1):
            c = ws.cell(row=row, column=col, value=val)
            c.fill = _fill(bg)
            c.font = _font(bold=True, color=fg, size=size)
            c.alignment = _align("center", "center")
            c.border = _border()

    def _row(ws, row: int, data: list, bg=None, bold=False, num_cols: set = None):
        for col, val in enumerate(data, 1):
            c = ws.cell(row=row, column=col, value=val)
            if bg:
                c.fill = _fill(bg)
            c.font = _font(bold=bold, size=9)
            c.alignment = _align("left", "center", wrap=True)
            c.border = _border()
            if num_cols and col in num_cols:
                c.alignment = _align("right", "center")

    wb = Workbook()
    wb.remove(wb.active)

    # ─── Feuille 1 : Synthèse ────────────────────────────────────
    ws1 = wb.create_sheet("Synthèse")
    ws1.column_dimensions["A"].width = 30
    ws1.column_dimensions["B"].width = 28

    # Titre
    ws1.merge_cells("A1:B1")
    c = ws1["A1"]
    c.value = f"RAPPORT DE GESTION — {payload['station_nom']}"
    c.fill = _fill(NAVY_H)
    c.font = _font(bold=True, color=WHITE, size=14)
    c.alignment = _align("center", "center")
    ws1.row_dimensions[1].height = 28

    ws1.merge_cells("A2:B2")
    c2 = ws1["A2"]
    c2.value = f"{payload['date_debut']}  →  {payload['date_fin']}  ({payload['nb_jours']} jours)"
    c2.fill = _fill(AMBER_H)
    c2.font = _font(bold=True, color=NAVY_H, size=11)
    c2.alignment = _align("center", "center")
    ws1.row_dimensions[2].height = 20

    kpis = [
        ("Volume total vendu", f"{stats['total_quantite']:,.3f} gal"),
        ("Revenu total", f"{stats['total_montant']:,.2f} G"),
        ("Relevés analysés", str(stats["nb_releves"])),
        ("Jours actifs", str(stats["nb_jours_couverts"])),
        ("Variation vs période préc.", _fmt_pct(narrative["var_pct"])),
        ("Anomalies détectées", str(len(anomalies))),
        ("Généré le", payload["date_generation"]),
    ]
    for i, (k, v) in enumerate(kpis, 3):
        ws1.cell(i, 1, k).fill = _fill(GREY_H)
        ws1.cell(i, 1).font = _font(bold=True, size=9)
        ws1.cell(i, 1).alignment = _align()
        ws1.cell(i, 1).border = _border()
        ws1.cell(i, 2, v).font = _font(size=9)
        ws1.cell(i, 2).alignment = _align()
        ws1.cell(i, 2).border = _border()

    # Résumé narratif
    row_narr = len(kpis) + 4
    ws1.merge_cells(f"A{row_narr}:B{row_narr}")
    ws1.cell(row_narr, 1, "Résumé").fill = _fill(TEAL_H)
    ws1.cell(row_narr, 1).font = _font(bold=True, color=WHITE, size=10)
    ws1.cell(row_narr, 1).alignment = _align("center", "center")

    row_narr += 1
    ws1.merge_cells(f"A{row_narr}:B{row_narr+3}")
    nc = ws1.cell(row_narr, 1, narrative["intro_kpis"])
    nc.alignment = _align("left", "top", wrap=True)
    nc.font = _font(size=9)
    ws1.row_dimensions[row_narr].height = 60

    # Conclusion
    row_conc = row_narr + 5
    ws1.merge_cells(f"A{row_conc}:B{row_conc}")
    ws1.cell(row_conc, 1, "Conclusion").fill = _fill(NAVY_H)
    ws1.cell(row_conc, 1).font = _font(bold=True, color=WHITE)
    ws1.cell(row_conc, 1).alignment = _align("center", "center")

    row_conc += 1
    ws1.merge_cells(f"A{row_conc}:B{row_conc+3}")
    cc = ws1.cell(row_conc, 1, narrative["conclusion"])
    cc.alignment = _align("left", "top", wrap=True)
    cc.font = _font(size=9)
    ws1.row_dimensions[row_conc].height = 60

    # Recommandations
    row_rec = row_conc + 5
    ws1.merge_cells(f"A{row_rec}:B{row_rec}")
    ws1.cell(row_rec, 1, "Recommandations").fill = _fill(AMBER_H)
    ws1.cell(row_rec, 1).font = _font(bold=True, color=NAVY_H)
    ws1.cell(row_rec, 1).alignment = _align("center", "center")
    for i, rec in enumerate(narrative["recommandations"], row_rec + 1):
        ws1.merge_cells(f"A{i}:B{i}")
        c_rec = ws1.cell(i, 1, f"{i - row_rec}. {rec}")
        c_rec.fill = _fill(GREY_H)
        c_rec.font = _font(size=9)
        c_rec.alignment = _align(wrap=True)
        c_rec.border = _border()
        ws1.row_dimensions[i].height = 30

    # ─── Feuille 2 : Ventes détaillées ───────────────────────────
    ws2 = wb.create_sheet("Ventes détaillées")
    for col, width in zip("ABCDE", [15, 18, 18, 18, 15]):
        ws2.column_dimensions[get_column_letter(col if isinstance(col, int) else ord(col) - 64)].width = width

    ws2.freeze_panes = "A2"
    _hdr(ws2, 1, ["Date", "Volume (gal)", "Revenu (G)", "Période Matin (G)", "Période AM (G)"])

    dates_tri = sorted(serie.keys())
    for i, d in enumerate(dates_tri, 2):
        v = serie[d]
        _row(ws2, i,
             [d, v["quantite"], v["montant"], "", ""],
             bg=GREY_H if i % 2 == 0 else WHITE,
             num_cols={2, 3})

    # Total
    last_row = len(dates_tri) + 2
    _hdr(ws2, last_row,
         ["TOTAL",
          round(sum(v["quantite"] for v in serie.values()), 3),
          round(sum(v["montant"] for v in serie.values()), 2),
          "", ""],
         bg=AMBER_H, fg=NAVY_H)

    # Graphique barres natif Excel
    if len(dates_tri) > 0:
        chart = BarChart()
        chart.type = "col"
        chart.title = "Ventes journalières"
        chart.y_axis.title = "Revenu (G)"
        chart.x_axis.title = "Date"
        chart.shape = 4
        data_ref = Reference(ws2, min_col=3, min_row=1, max_row=len(dates_tri) + 1)
        cats_ref = Reference(ws2, min_col=1, min_row=2, max_row=len(dates_tri) + 1)
        chart.add_data(data_ref, titles_from_data=True)
        chart.set_categories(cats_ref)
        chart.series[0].graphicalProperties.solidFill = AMBER_H
        chart.width = 20
        chart.height = 12
        ws2.add_chart(chart, f"A{last_row + 2}")

    # ─── Feuille 3 : Répartition ──────────────────────────────────
    ws3 = wb.create_sheet("Répartition")
    ws3.freeze_panes = "A2"

    ws3.column_dimensions["A"].width = 20
    ws3.column_dimensions["B"].width = 20
    ws3.column_dimensions["C"].width = 18
    ws3.column_dimensions["D"].width = 18
    ws3.column_dimensions["E"].width = 12

    _hdr(ws3, 1, ["Produit/Pompe", "Type", "Volume (gal)", "Revenu (G)", "% revenu"])
    row3 = 2
    for nom, d in sorted(stats["par_produit"].items(), key=lambda x: x[1]["montant"], reverse=True):
        pct = round(d["montant"] / stats["total_montant"] * 100, 1) if stats["total_montant"] > 0 else 0.0
        _row(ws3, row3, [nom, "Produit", d["quantite"], d["montant"], f"{pct:.1f} %"],
             bg=TEAL_L if row3 % 2 == 0 else WHITE, num_cols={3, 4})
        row3 += 1

    ws3.cell(row3, 1, "").fill = _fill(GREY_H)
    row3 += 1

    for nom, d in sorted(stats["par_pompe"].items(), key=lambda x: x[1]["montant"], reverse=True):
        pct = round(d["montant"] / stats["total_montant"] * 100, 1) if stats["total_montant"] > 0 else 0.0
        _row(ws3, row3, [nom, f"Pompe ({d['produit']})", d["quantite"], d["montant"], f"{pct:.1f} %"],
             bg=GREY_H if row3 % 2 == 0 else WHITE, num_cols={3, 4})
        row3 += 1

    # Camembert natif Excel (par produit)
    if len(stats["par_produit"]) > 1:
        pie = PieChart()
        pie.title = "Revenu par produit"
        labels_ref = Reference(ws3, min_col=1, min_row=2, max_row=1 + len(stats["par_produit"]))
        data_pie = Reference(ws3, min_col=4, min_row=1, max_row=1 + len(stats["par_produit"]))
        pie.add_data(data_pie, titles_from_data=True)
        pie.set_categories(labels_ref)
        pie.width = 14
        pie.height = 12
        ws3.add_chart(pie, f"G2")

    # ─── Feuille 4 : Anomalies ────────────────────────────────────
    ws4 = wb.create_sheet("Anomalies")
    ws4.freeze_panes = "A2"
    ws4.column_dimensions["A"].width = 22
    ws4.column_dimensions["B"].width = 12
    ws4.column_dimensions["C"].width = 20
    ws4.column_dimensions["D"].width = 12
    ws4.column_dimensions["E"].width = 14
    ws4.column_dimensions["F"].width = 55

    _hdr(ws4, 1, ["Type", "Gravité", "Pompe/Produit", "Date", "Période", "Message"])
    for i, a in enumerate(anomalies, 2):
        bg = REDL_H if a.get("gravite") == "erreur" else "FEF3C7"
        _row(ws4, i, [
            a.get("type", "—"),
            a.get("gravite", "—"),
            a.get("pompe_nom", a.get("produit_nom", "—")),
            a.get("date", "—"),
            a.get("periode", "—"),
            (a.get("message", "") or "")[:200],
        ], bg=bg, num_cols=set())
        ws4.row_dimensions[i].height = 28

    if not anomalies:
        ws4.merge_cells("A2:F2")
        ws4["A2"].value = "Aucune anomalie détectée — données de qualité."
        ws4["A2"].font = _font(bold=True, color=TEAL_H)
        ws4["A2"].alignment = _align("center")

    # ─── Feuille 5 : Stock & Rentabilité ─────────────────────────
    ws5 = wb.create_sheet("Stock & Rentabilité")
    ws5.column_dimensions["A"].width = 20

    # Stock
    _hdr(ws5, 1, ["Produit", "Stock (gal)", "Livrés (gal)", "Vendus (gal)", "Jours restants", "Alerte"])
    for col, w in zip("BCDEF", [14, 14, 14, 14, 10]):
        ws5.column_dimensions[get_column_letter(ord(col) - 64)].width = w
    for i, s in enumerate(stocks, 2):
        alerte = "⚠ BAS" if s["alerte_bas"] else "OK"
        bg = REDL_H if s["alerte_bas"] else (GREY_H if i % 2 == 0 else WHITE)
        _row(ws5, i, [
            s["produit_nom"],
            s["gallons_restants"],
            s["gallons_livres"],
            s["gallons_vendus"],
            s["jours_de_stock"] if s["jours_de_stock"] is not None else "—",
            alerte,
        ], bg=bg, num_cols={2, 3, 4, 5})

    # Séparateur
    row5 = len(stocks) + 3
    ws5.merge_cells(f"A{row5}:F{row5}")
    ws5.cell(row5, 1, "RENTABILITÉ — Coût Moyen Pondéré (WAC)").fill = _fill(NAVY_H)
    ws5.cell(row5, 1).font = _font(bold=True, color=WHITE)
    ws5.cell(row5, 1).alignment = _align("center")

    row5 += 1
    _hdr(ws5, row5, ["Produit", "Vendus (gal)", "Revenu (G)", "COGS (G)", "Bénéfice (G)", "Marge %"])
    for p in rentab.get("produits", []):
        row5 += 1
        _row(ws5, row5, [
            p["produit_nom"],
            p["gallons_vendus"],
            p["revenu_total"],
            p["cogs_total"] if p["cogs_total"] is not None else "—",
            p["benefice"] if p["benefice"] is not None else "—",
            f"{p['marge_pct']:.2f} %" if p["marge_pct"] is not None else "—",
        ], bg=GREY_H if row5 % 2 == 0 else WHITE, num_cols={2, 3, 4, 5})

    t_tot = rentab.get("total", {})
    row5 += 1
    _hdr(ws5, row5, [
        "TOTAL",
        t_tot.get("gallons_vendus", 0),
        t_tot.get("revenu_total", 0),
        t_tot.get("cogs_total", 0) if t_tot.get("cogs_total") is not None else "—",
        t_tot.get("benefice", 0) if t_tot.get("benefice") is not None else "—",
        f"{t_tot.get('marge_pct', 0):.2f} %" if t_tot.get("marge_pct") is not None else "—",
    ], bg=AMBER_H, fg=NAVY_H)

    # Auto-filter sur feuille anomalies
    if anomalies:
        ws4.auto_filter.ref = f"A1:F{len(anomalies)+1}"

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()
