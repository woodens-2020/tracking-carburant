"""
Rendeurs pour le document d'audit IA — PDF et DOCX (v2).

Nouveautés v2 :
  • Bandeau KPI 8 métriques (2 lignes × 4 colonnes) en têtes colorées
  • Colorisation automatique des montants (G → amber), volumes (gal → teal),
    pourcentages (% → violet) dans les paragraphes PDF
  • Runs colorés pour les **bold** DOCX selon le type de donnée
  • Titre de chaque section surligné d'une barre de couleur
"""
from __future__ import annotations

import re
from datetime import date
from io import BytesIO


# ══════════════════════════════════════════════════════════════════════════════
# Parser Markdown → blocs structurés
# ══════════════════════════════════════════════════════════════════════════════

def _parse(text: str) -> list[dict]:
    blocks = []
    for line in text.split("\n"):
        s = line.strip()
        if not s:
            blocks.append({"t": "space"})
        elif s.startswith("#### "):
            blocks.append({"t": "h4", "c": s[5:]})
        elif s.startswith("### "):
            blocks.append({"t": "h3", "c": s[4:]})
        elif s.startswith("## "):
            blocks.append({"t": "h2", "c": s[3:]})
        elif s.startswith("# "):
            blocks.append({"t": "h1", "c": s[2:]})
        elif re.match(r"^---+$", s):
            blocks.append({"t": "hr"})
        elif re.match(r"^(\d+\.|[•\-\*])\s", s):
            clean = re.sub(r"^(\d+\.\s*|[•\-\*]\s*)", "", s)
            blocks.append({"t": "li", "c": clean})
        elif s.startswith("*") and s.endswith("*") and not s.startswith("**"):
            blocks.append({"t": "em", "c": s.strip("*")})
        else:
            blocks.append({"t": "p", "c": s})
    return blocks


def _strip_bold(text: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


# ══════════════════════════════════════════════════════════════════════════════
# Colorisation PDF : nombres → balises reportlab
# ══════════════════════════════════════════════════════════════════════════════

_AMBER_HEX  = "#f7a93b"
_TEAL_HEX   = "#3fb6a8"
_VIOLET_HEX = "#7c3aed"
_RED_HEX    = "#e53e3e"

_RE_G   = re.compile(r"(\d[\d ]*(?:[.,]\d+)?\s*G\b)")
_RE_GAL = re.compile(r"(\d[\d ]*(?:[.,]\d+)?\s*gal\b)")
_RE_PCT = re.compile(r"(\d+(?:[.,]\d+)?\s*%)")
_RE_TAG = re.compile(r"<[^>]+>")


def _colorize_segment(seg: str) -> str:
    """Colorise montants/volumes/% dans un segment de texte pur (sans balises HTML)."""
    seg = _RE_G.sub(rf'<font color="{_AMBER_HEX}">\1</font>', seg)
    seg = _RE_GAL.sub(rf'<font color="{_TEAL_HEX}">\1</font>', seg)
    seg = _RE_PCT.sub(rf'<font color="{_VIOLET_HEX}">\1</font>', seg)
    return seg


def _enrich_para(text: str) -> str:
    """
    1. Convertit **bold** → <b>bold</b>
    2. Colorise les montants (G), volumes (gal) et pourcentages
       uniquement dans les segments de texte (pas dans les attrs de balises).
    """
    # Bold markdown
    result = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    # Colorise uniquement les segments non-tag
    parts = []
    pos = 0
    for m in _RE_TAG.finditer(result):
        parts.append(_colorize_segment(result[pos:m.start()]))
        parts.append(m.group())
        pos = m.end()
    parts.append(_colorize_segment(result[pos:]))
    return "".join(parts)


# ══════════════════════════════════════════════════════════════════════════════
# Rendu PDF
# ══════════════════════════════════════════════════════════════════════════════

def render_audit_pdf(
    text: str,
    date_debut: date,
    date_fin: date,
    kpi_data: dict | None = None,
) -> bytes:
    from reportlab.lib          import colors
    from reportlab.lib.units    import cm
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles  import ParagraphStyle
    from reportlab.lib.enums   import TA_LEFT, TA_CENTER, TA_JUSTIFY
    from reportlab.platypus    import (
        SimpleDocTemplate, Paragraph, Spacer,
        HRFlowable, PageBreak, KeepTogether, Table, TableStyle,
    )

    NAVY   = colors.HexColor("#0f1e35")
    AMBER  = colors.HexColor("#f7a93b")
    VIOLET = colors.HexColor("#7c3aed")
    TEAL   = colors.HexColor("#3fb6a8")
    GRAY   = colors.HexColor("#555555")
    LGRAY  = colors.HexColor("#888888")
    WHITE  = colors.white

    # Card colors for KPI banner
    C_NAVY = colors.HexColor("#0f1e35")
    C_VIOL = colors.HexColor("#5a1eba")
    C_TEAL = colors.HexColor("#1a8a7c")
    C_AMBE = colors.HexColor("#b87000")
    C_DARK = colors.HexColor("#333344")
    C_FOREST = colors.HexColor("#1a6b40")

    def _s(name, **kw):
        return ParagraphStyle(name, **kw)

    s_meta   = _s("meta",   fontName="Helvetica-Oblique",    fontSize=9,  textColor=LGRAY, spaceAfter=3,  alignment=TA_CENTER)
    s_cover1 = _s("cov1",   fontName="Helvetica-Bold",       fontSize=22, textColor=NAVY,  spaceAfter=6,  alignment=TA_CENTER)
    s_cover2 = _s("cov2",   fontName="Helvetica-Bold",       fontSize=13, textColor=VIOLET,spaceAfter=4,  alignment=TA_CENTER)
    s_h1     = _s("H1",     fontName="Helvetica-Bold",       fontSize=14, textColor=WHITE,
                  backColor=NAVY, spaceBefore=18, spaceAfter=10, leftIndent=8, leading=22)
    s_h2     = _s("H2",     fontName="Helvetica-Bold",       fontSize=12, textColor=NAVY,
                  spaceBefore=14, spaceAfter=5)
    s_h3     = _s("H3",     fontName="Helvetica-Bold",       fontSize=11, textColor=VIOLET,
                  spaceBefore=10, spaceAfter=4)
    s_h4     = _s("H4",     fontName="Helvetica-BoldOblique",fontSize=10, textColor=TEAL,
                  spaceBefore=7,  spaceAfter=3)
    s_para   = _s("para",   fontName="Helvetica",            fontSize=10, textColor=colors.black,
                  leading=16, spaceAfter=6, alignment=TA_JUSTIFY)
    s_li     = _s("li",     fontName="Helvetica",            fontSize=10, textColor=colors.black,
                  leading=15, spaceAfter=4, leftIndent=22, bulletIndent=8)
    s_em     = _s("em",     fontName="Helvetica-Oblique",    fontSize=9,  textColor=LGRAY,
                  spaceAfter=4, alignment=TA_CENTER)
    # KPI banner styles
    s_klbl   = _s("klbl",   fontName="Helvetica",            fontSize=7,  textColor=WHITE,  alignment=TA_CENTER, spaceAfter=1)
    s_kval   = _s("kval",   fontName="Helvetica-Bold",       fontSize=13, textColor=WHITE,  alignment=TA_CENTER, leading=16)
    s_ksub   = _s("ksub",   fontName="Helvetica-Oblique",    fontSize=7,  textColor=colors.HexColor("#cccccc"), alignment=TA_CENTER)

    # ── Helpers formatage KPI ──
    def _kg(v):
        if v is None: return "—"
        if abs(v) >= 1_000_000:
            return f"{v/1_000_000:.2f} M G"
        return f"{v:,.0f} G".replace(",", " ")

    def _kpct(v):
        if v is None: return "—"
        sign = "▲ " if v >= 0 else "▼ "
        return f"{sign}{abs(v):.1f} %"

    def _kpct_plain(v):
        if v is None: return "—"
        return f"{v:.1f} %"

    def _kg2(v):
        if v is None: return "—"
        return f"{v:.2f} G/gal"

    # ── Bandeau KPI ──────────────────────────────────────────────────────────
    def _build_kpi_banner(kd: dict) -> list:
        """Construit 2 rangées de 4 tuiles KPI chacune."""
        rows_data = [
            [
                ("REVENU TOTAL",        _kg(kd.get("revenu_total")),      C_NAVY),
                ("BÉNÉFICE BRUT",       _kg(kd.get("benefice_brut")),     C_VIOL),
                ("MARGE BRUTE",         _kpct_plain(kd.get("marge_pct")), C_TEAL),
                ("REV. MOYEN / JOUR",   _kg(kd.get("revenu_par_jour")),   C_AMBE),
            ],
            [
                ("PRIX VENTE MOY.",     _kg2(kd.get("prix_vente_moyen")), C_DARK),
                ("TAUX COUVERTURE",     _kpct_plain(kd.get("taux_couverture")), C_FOREST),
                ("VAR. VS PRÉCÉDENT",   _kpct(kd.get("var_revenu_pct")),  C_VIOL),
                ("ANOMALIES DÉTECTÉES", f"{kd.get('nb_anomalies', 0)} total  ·  {kd.get('nb_erreurs', 0)} erreurs", C_DARK),
            ],
        ]

        flowables = []
        col_w = (A4[0] - 4.4 * cm) / 4

        for row_cards in rows_data:
            label_cells = [Paragraph(c[0], s_klbl) for c in row_cards]
            value_cells = [Paragraph(c[1], s_kval) for c in row_cards]
            bgcolors    = [c[2] for c in row_cards]

            table_data = [label_cells, value_cells]
            tbl = Table(table_data, colWidths=[col_w] * 4, rowHeights=[14, 28])
            ts  = TableStyle([
                ("ALIGN",        (0, 0), (-1, -1), "CENTER"),
                ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING",   (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING",(0, 0), (-1, -1), 5),
                ("LEFTPADDING",  (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("GRID",         (0, 0), (-1, -1), 0.4, WHITE),
            ])
            for col_idx, bg in enumerate(bgcolors):
                ts.add("BACKGROUND", (col_idx, 0), (col_idx, 1), bg)
            tbl.setStyle(ts)
            flowables.append(tbl)
            flowables.append(Spacer(1, 3))

        return flowables

    # ── Construction du document ─────────────────────────────────────────────
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2.2 * cm, rightMargin=2.2 * cm,
        topMargin=2.5 * cm,  bottomMargin=2.2 * cm,
    )

    story = []

    # Page de couverture
    story.append(Spacer(1, 0.8 * cm))
    story.append(Paragraph("DOCUMENT D'AUDIT CONFIDENTIEL", s_meta))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph("PétroSync · Station Carburant", s_cover1))
    story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph(
        f"Période analysée : {date_debut.strftime('%d/%m/%Y')} – {date_fin.strftime('%d/%m/%Y')}",
        s_cover2,
    ))
    story.append(Spacer(1, 0.15 * cm))
    story.append(Paragraph(
        f"Généré le {date.today().strftime('%d/%m/%Y')} par Intelligence Artificielle",
        s_meta,
    ))
    story.append(HRFlowable(width="100%", thickness=2.5, color=AMBER, spaceAfter=12, spaceBefore=10))

    # Bandeau KPI
    if kpi_data:
        story.append(Paragraph("TABLEAU DE BORD — INDICATEURS CLÉS", s_meta))
        story.append(Spacer(1, 4))
        story.extend(_build_kpi_banner(kpi_data))
        story.append(HRFlowable(width="100%", thickness=1, color=LGRAY, spaceAfter=6, spaceBefore=6))

    story.append(Spacer(1, 0.3 * cm))

    # Contenu IA
    blocks = _parse(text)
    skip_h1 = True

    for b in blocks:
        t = b["t"]

        if t == "space":
            story.append(Spacer(1, 4))

        elif t == "hr":
            story.append(HRFlowable(width="100%", thickness=0.6, color=GRAY, spaceAfter=8, spaceBefore=6))

        elif t == "h1":
            if skip_h1:
                skip_h1 = False
                continue
            story.append(PageBreak())
            story.append(Paragraph(_strip_bold(b["c"]), s_h1))

        elif t == "h2":
            story.append(KeepTogether([
                Paragraph(_strip_bold(b["c"]), s_h2),
                HRFlowable(width="100%", thickness=1, color=AMBER, spaceAfter=4),
            ]))

        elif t == "h3":
            story.append(Paragraph(_strip_bold(b["c"]), s_h3))

        elif t == "h4":
            story.append(Paragraph(_strip_bold(b["c"]), s_h4))

        elif t == "li":
            story.append(Paragraph(f"• {_enrich_para(b['c'])}", s_li))

        elif t == "em":
            story.append(Paragraph(f"<i>{b['c']}</i>", s_em))

        elif t == "p":
            if not b["c"]:
                continue
            story.append(Paragraph(_enrich_para(b["c"]), s_para))

    doc.build(story)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# Rendu DOCX
# ══════════════════════════════════════════════════════════════════════════════

def render_audit_docx(
    text: str,
    date_debut: date,
    date_fin: date,
    kpi_data: dict | None = None,
) -> bytes:
    from docx           import Document
    from docx.shared    import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns   import qn
    from docx.oxml      import OxmlElement

    NAVY   = RGBColor(0x0f, 0x1e, 0x35)
    AMBER  = RGBColor(0xf7, 0xa9, 0x3b)
    VIOLET = RGBColor(0x7c, 0x3a, 0xed)
    TEAL   = RGBColor(0x3f, 0xb6, 0xa8)
    GRAY   = RGBColor(0x88, 0x88, 0x88)
    WHITE  = RGBColor(0xff, 0xff, 0xff)

    doc = Document()
    for section in doc.sections:
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.2)

    def _set_color(run, rgb):
        run.font.color.rgb = rgb

    def _detect_color(text: str):
        """Retourne la couleur appropriée selon le contenu du segment bold."""
        if re.search(r"\d+.*G\b", text):
            return AMBER
        if re.search(r"\d+.*gal\b", text):
            return TEAL
        if re.search(r"\d+.*%", text):
            return VIOLET
        return None

    def _add_run_md(para, content):
        """Texte avec **gras** colorisé selon le contexte (montant/volume/%)."""
        parts = re.split(r"(\*\*[^*]+\*\*)", content)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                inner = part[2:-2]
                r = para.add_run(inner)
                r.bold = True
                col = _detect_color(inner)
                if col:
                    _set_color(r, col)
            elif part:
                para.add_run(part)

    def _cell_shading(cell, hex_color: str):
        """Applique un arrière-plan coloré à une cellule Word."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color.lstrip("#"))
        tcPr.append(shd)

    def _kg(v):
        if v is None: return "N/D"
        if abs(v) >= 1_000_000:
            return f"{v/1_000_000:.2f} M G"
        return f"{v:,.0f} G".replace(",", " ")

    def _kpct(v, sign=False):
        if v is None: return "N/D"
        if sign:
            arrow = "▲" if v >= 0 else "▼"
            return f"{arrow} {abs(v):.1f} %"
        return f"{v:.1f} %"

    def _kg2(v):
        if v is None: return "N/D"
        return f"{v:.2f} G/gal"

    # ── Couverture ───────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DOCUMENT D'AUDIT CONFIDENTIEL")
    r.font.size = Pt(9); r.font.italic = True; _set_color(r, GRAY)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("PétroSync · Station Carburant")
    r2.font.size = Pt(22); r2.bold = True; _set_color(r2, NAVY)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(
        f"Période : {date_debut.strftime('%d/%m/%Y')} – {date_fin.strftime('%d/%m/%Y')}"
    )
    r3.font.size = Pt(14); r3.bold = True; _set_color(r3, VIOLET)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(f"Généré le {date.today().strftime('%d/%m/%Y')} par Intelligence Artificielle")
    r4.font.size = Pt(9); r4.italic = True; _set_color(r4, GRAY)

    # ── Tableau KPI DOCX ─────────────────────────────────────────────────────
    if kpi_data:
        doc.add_paragraph()
        hdr = doc.add_paragraph()
        hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_hdr = hdr.add_run("TABLEAU DE BORD — INDICATEURS CLÉS")
        r_hdr.bold = True; r_hdr.font.size = Pt(9); _set_color(r_hdr, GRAY)

        kpi_rows = [
            [
                ("REVENU TOTAL",       _kg(kpi_data.get("revenu_total")),             "0f1e35"),
                ("BÉNÉFICE BRUT",      _kg(kpi_data.get("benefice_brut")),            "5a1eba"),
                ("MARGE BRUTE",        _kpct(kpi_data.get("marge_pct")),              "1a8a7c"),
                ("REV. / JOUR",        _kg(kpi_data.get("revenu_par_jour")),          "b87000"),
            ],
            [
                ("PRIX VENTE MOY.",    _kg2(kpi_data.get("prix_vente_moyen")),        "333344"),
                ("TAUX COUVERTURE",    _kpct(kpi_data.get("taux_couverture")),        "1a6b40"),
                ("VAR. REVENU",        _kpct(kpi_data.get("var_revenu_pct"), sign=True), "5a1eba"),
                ("ANOMALIES",          f"{kpi_data.get('nb_anomalies', 0)} / {kpi_data.get('nb_erreurs', 0)} err.", "7a3010"),
            ],
        ]

        for row_cards in kpi_rows:
            tbl = doc.add_table(rows=2, cols=4)
            tbl.style = "Table Grid"
            for col_idx, (label, value, bg_hex) in enumerate(row_cards):
                label_cell = tbl.rows[0].cells[col_idx]
                value_cell = tbl.rows[1].cells[col_idx]
                _cell_shading(label_cell, bg_hex)
                _cell_shading(value_cell, bg_hex)
                lp = label_cell.paragraphs[0]
                lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                lr = lp.add_run(label)
                lr.font.size = Pt(7); lr.bold = True; _set_color(lr, WHITE)
                vp = value_cell.paragraphs[0]
                vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                vr = vp.add_run(value)
                vr.font.size = Pt(12); vr.bold = True; _set_color(vr, WHITE)
            doc.add_paragraph()

    doc.add_paragraph()

    # ── Contenu IA ───────────────────────────────────────────────────────────
    blocks = _parse(text)
    skip_h1 = True

    for b in blocks:
        t = b["t"]

        if t == "space":
            continue

        elif t == "hr":
            p = doc.add_paragraph()
            r = p.add_run("─" * 65)
            r.font.size = Pt(7); _set_color(r, GRAY)

        elif t == "h1":
            if skip_h1:
                skip_h1 = False
                continue
            h = doc.add_heading(_strip_bold(b["c"]), level=1)
            for run in h.runs:
                _set_color(run, NAVY)

        elif t == "h2":
            h = doc.add_heading(_strip_bold(b["c"]), level=2)
            for run in h.runs:
                _set_color(run, NAVY)

        elif t == "h3":
            h = doc.add_heading(_strip_bold(b["c"]), level=3)
            for run in h.runs:
                _set_color(run, VIOLET)

        elif t == "h4":
            h = doc.add_heading(_strip_bold(b["c"]), level=4)
            for run in h.runs:
                _set_color(run, TEAL)

        elif t == "li":
            p = doc.add_paragraph(style="List Bullet")
            _add_run_md(p, b["c"])

        elif t == "em":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(b["c"])
            r.font.size = Pt(9); r.italic = True; _set_color(r, GRAY)

        elif t == "p":
            if not b["c"]:
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_run_md(p, b["c"])

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
